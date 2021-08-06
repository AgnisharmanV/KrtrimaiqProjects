import boto3
from docx import Document
import numpy as np
import sys
import time


class AWSRef:
    def __init__(self, bucket, region):
        self.s3 = boto3.client("s3")
        self.textract = boto3.client("textract")
        self.bucket = bucket
        self.create_bucket()
        self.region = region

    def upload_obj(self, bytes_obj, name, acl):
        self.s3.put_object(
            Body=bytes_obj,
            ACL=acl,
            Bucket=self.bucket,
            Key=name
        )

    def cleanup(self, key):
        self.s3.delete_object(
            Bucket=self.bucket,
            Key=key
        )

    def get_url(self, key):
        url = f'https://{self.bucket}.s3.{self.region}.amazonaws.com/{key}'
        return url

    def create_bucket(self):
        """Creates bucket if not present"""
        bucket_names = []
        for bucket_info in self.s3.list_buckets()['Buckets']:
            bucket_names.append(bucket_info['Name'])

        if self.bucket not in bucket_names:
            self.s3.create_bucket(
                ACL='private',
                Bucket=self.bucket,
                CreateBucketConfiguration={
                    'LocationConstraint': 'ap-south-1'
                }
            )

    def analyze_doc(self, bytes_obj):
        response = self.textract.analyze_document(
            Document={'Bytes': bytes_obj},
            FeatureTypes=['TABLES']
        )
        return response['Blocks']

    def start_analysis(self, key):
        response = self.textract.start_document_analysis(
            DocumentLocation={
                'S3Object': {
                    'Bucket': self.bucket,
                    'Name': key
                }
            },
            FeatureTypes=['FORMS']
        )
        return response['JobId']

    def check_progress(self, jobId, nextToken, timelim=300):
        while timelim>0:
            if nextToken:
                response = self.textract.get_document_analysis(
                    JobId=jobId,
                    MaxResults=1000,
                    NextToken=nextToken
                )
            else:
                response = self.textract.get_document_analysis(
                    JobId=jobId,
                    MaxResults=1000,
                )

            if response['JobStatus'] == 'SUCCEEDED':
                print('Document analysis succeeded.')
                return response

            elif response['JobStatus'] == 'FAILED':
                print('Document analysis failed.')
                sys.exit(0)

            else:
                print('Document analysis in progress...')
                time.sleep(20)
                timelim -= 20

            if timelim <= 0:
                print('Timeout')
                sys.exit(0)


class Docx:
    def __init__(self):
        self.doc = Document()

    def write_paragraph(self, para, extract=True):
        if extract:
            self.doc.add_paragraph(para)

    def write_table(self, arr, extract=False):
        if extract:
            table = self.doc.add_table(rows=arr.shape[0], cols=arr.shape[1])
            for row in range(arr.shape[0]):
                row_cells = table.rows[row].cells
                for col in range(arr.shape[1]):
                    row_cells[col].text = arr[row][col]

    def save_doc(self, file_obj):
        self.doc.save(file_obj)


def part_of_table(line_block, table_block, blocks_map):
    """Checks whether block in part of table"""
    line_words = []
    if 'Relationships' in line_block.keys():
        for relationship in line_block['Relationships']:
            for idx in relationship['Ids']:
                if blocks_map[idx]['BlockType'] == 'WORD':
                    line_words.append(idx)
    else:
        return False

    table_words = []
    if 'Relationships' in table_block.keys():
        for relationship in table_block['Relationships']:
            for idx in relationship['Ids']:
                if blocks_map[idx]['BlockType'] == 'WORD':
                    table_words.append(idx)
                if blocks_map[idx]['BlockType'] in ['LINE', 'CELL']:
                    sub_block = blocks_map[idx]

                    if 'Relationships' in sub_block.keys():
                        for rela in sub_block['Relationships']:
                            for i in rela['Ids']:
                                if blocks_map[i]['BlockType'] == 'WORD':
                                    table_words.append(i)
    else:
        return False

    flag = False
    for l in line_words:
        if l in table_words:
            return True
    return flag


def get_text(result, blocks_map):
    text = ''
    if 'Relationships' in result:
        for relationship in result['Relationships']:
            if relationship['Type'] == 'CHILD':
                for child_id in relationship['Ids']:
                    word = blocks_map[child_id]
                    if word['BlockType'] == 'WORD':
                        text += word['Text'] + ' '
                    if word['BlockType'] == 'SELECTION_ELEMENT':
                        if word['SelectionStatus'] =='SELECTED':
                            text +=  'X '
    return text


def get_rows_columns_map(table_result, blocks_map):
    rows = {}
    for relationship in table_result['Relationships']:
        if relationship['Type'] == 'CHILD':
            for child_id in relationship['Ids']:
                cell = blocks_map[child_id]
                if cell['BlockType'] == 'CELL':
                    row_index = cell['RowIndex']
                    col_index = cell['ColumnIndex']
                    if row_index not in rows:
                        # create new row
                        rows[row_index] = {}

                    # get the text value
                    rows[row_index][col_index] = get_text(cell, blocks_map)
    return rows


def make_table(table_block, blocks_map):
    table = get_rows_columns_map(table_block, blocks_map)
    n_rows = len(table.keys())
    n_cols = len(table[list(table.keys())[0]].keys())
    arr_row = []
    arr = []
    for r in range(n_rows):
        for c in range(n_cols):
            arr_row.append(table[r+1][c+1])
        arr.append(arr_row)
        arr_row = []
    return np.array(arr)


def write_to_docx(response, doc_instance, tags=None):
    blocks_map = {}
    table_blocks = []
    line_blocks = []
    for block in response:
        blocks_map[block['Id']] = block
        if block['BlockType'] == "TABLE":
            table_blocks.append(block)
        if block['BlockType'] == 'LINE':
            line_blocks.append(block)

    text = ''
    writing = False
    in_table = False
    extract = True
    if tags:
        extract = False

    for b, block in enumerate(line_blocks):
        if tags:
            if tags[0] in block['Text'].lower():
                extract = True
            if tags[1].lower() in block['Text'].lower():
                extract = False

        if len(table_blocks) > 0:
            if not in_table:
                in_table = part_of_table(block, table_blocks[0], blocks_map)
                writing = True
            if in_table and writing:
                t = make_table(table_blocks[0], blocks_map)
                doc_instance.write_paragraph(text)
                text = ''
                doc_instance.write_table(t, extract)
                writing = False
            elif not in_table:
                if extract:
                    text = text + block['Text'] + '\n'
            elif in_table and (not writing):
                if b < (len(line_blocks) - 1 ):
                    if not part_of_table(line_blocks[b + 1], table_blocks[0], blocks_map):
                        table_blocks.pop(0)
                        in_table = False
        else:
            if extract:
                text = text + block['Text'] + '\n'
    doc_instance.write_paragraph(text)


def write_to_text(response, tags=None):
    line_blocks = []
    for block in response:
        if block['BlockType'] == 'LINE':
            line_blocks.append(block)

    text = ''
    extract = False
    for b, block in enumerate(line_blocks):
        if tags:
            if tags[0].lower() in block['Text'].lower():
                extract = True
            if tags[1].lower() in block['Text'].lower():
                extract = False
            if not extract:
                continue

        text = text + block['Text'] + '\n'

    return text
